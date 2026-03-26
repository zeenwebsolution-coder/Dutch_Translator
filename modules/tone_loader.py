"""
tone_loader.py — Parse the Dutch tone reference file.

Supported formats:
  .doc   → converted via LibreOffice, then read as .docx
  .docx  → python-docx (full paragraph text)
  .txt   → raw UTF-8
  .rtf   → stripped via LibreOffice → plain text
  .csv   → pandas, all cells joined
  .xlsx  → pandas, all sheets + cells joined
  .xls   → pandas (via xlrd), all sheets + cells joined
  .pdf   → pdfminer.six text extraction

Returns a single cleaned string of Dutch reference text.
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile

import pandas as pd
from modules.config import MAX_TONE_CHARS


class ToneLoadError(Exception):
    """Raised when the tone file cannot be parsed."""


# ── Dispatcher ────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".doc", ".docx", ".txt", ".rtf",
    ".csv", ".xlsx", ".xls", ".pdf",
}


def load_tone_file(uploaded_file) -> str:
    """
    Read the uploaded Streamlit file object and return cleaned Dutch text.

    Args:
        uploaded_file: Streamlit UploadedFile object.

    Returns:
        Whitespace-normalised Dutch reference text capped at MAX_TONE_CHARS chars.

    Raises:
        ToneLoadError: If the format is unsupported or the file is unreadable.
    """
    name = uploaded_file.name.lower()
    raw_bytes: bytes = uploaded_file.read()
    uploaded_file.seek(0)

    ext = _get_ext(name)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ToneLoadError(
            f"Unsupported tone file format: '{uploaded_file.name}'.\n"
            f"Accepted: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    try:
        if ext == ".txt":
            raw = _load_txt(raw_bytes)
        elif ext == ".docx":
            raw = _load_docx(raw_bytes)
        elif ext == ".doc":
            raw = _load_doc_via_libreoffice(raw_bytes)
        elif ext == ".rtf":
            raw = _load_rtf_via_libreoffice(raw_bytes)
        elif ext == ".csv":
            raw = _load_csv(raw_bytes)
        elif ext in (".xlsx", ".xls"):
            raw = _load_excel(raw_bytes)
        elif ext == ".pdf":
            raw = _load_pdf(raw_bytes)
        else:
            raise ToneLoadError(f"No handler for extension '{ext}'.")
    except ToneLoadError:
        raise
    except Exception as exc:
        raise ToneLoadError(
            f"Could not read '{uploaded_file.name}': {exc}"
        ) from exc

    cleaned = _clean(raw)
    if not cleaned:
        raise ToneLoadError(
            f"No text could be extracted from '{uploaded_file.name}'. "
            "Make sure the file contains readable Dutch text."
        )
    return cleaned[:MAX_TONE_CHARS]


# ── Format handlers ───────────────────────────────────────────────────────────

def _load_txt(data: bytes) -> str:
    # Try UTF-8 first, fall back to latin-1 (common in Dutch documents)
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _load_docx(data: bytes) -> str:
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return " ".join(paragraphs)


def _load_doc_via_libreoffice(data: bytes) -> str:
    """
    Convert legacy .doc to .docx using LibreOffice headless, then extract text.
    Falls back to mammoth if LibreOffice is unavailable.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")

    if soffice:
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "input.doc")
            with open(src, "wb") as f:
                f.write(data)
            try:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "docx",
                     "--outdir", tmpdir, src],
                    check=True,
                    timeout=60,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                out = os.path.join(tmpdir, "input.docx")
                if os.path.exists(out):
                    with open(out, "rb") as f:
                        return _load_docx(f.read())
            except (subprocess.SubprocessError, subprocess.TimeoutExpired):
                pass  # fall through to mammoth

    # Fallback: mammoth (works reasonably well on .doc too)
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(data))
        return result.value
    except Exception as exc:
        raise ToneLoadError(
            f"Could not convert .doc file. "
            f"Try saving it as .docx in Microsoft Word first. ({exc})"
        ) from exc


def _load_rtf_via_libreoffice(data: bytes) -> str:
    """Convert .rtf → .txt via LibreOffice headless."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise ToneLoadError(
            "LibreOffice is required to read .rtf files but was not found. "
            "Please convert your file to .txt or .docx."
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "input.rtf")
        with open(src, "wb") as f:
            f.write(data)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "txt:Text",
             "--outdir", tmpdir, src],
            check=True, timeout=60,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        out = os.path.join(tmpdir, "input.txt")
        if os.path.exists(out):
            with open(out, "rb") as f:
                return _load_txt(f.read())
    raise ToneLoadError("RTF conversion produced no output.")


def _load_csv(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            df = pd.read_csv(io.BytesIO(data), header=None, dtype=str, encoding=enc)
            return " ".join(df.fillna("").values.flatten().tolist())
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue
    raise ToneLoadError("Could not decode the CSV file. Try saving it as UTF-8.")


def _load_excel(data: bytes) -> str:
    sheets = pd.read_excel(
        io.BytesIO(data), sheet_name=None, header=None, dtype=str
    )
    parts: list[str] = []
    for df in sheets.values():
        parts.append(" ".join(df.fillna("").values.flatten().tolist()))
    return " ".join(parts)


def _load_pdf(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        out = io.StringIO()
        extract_text_to_fp(
            io.BytesIO(data), out,
            laparams=LAParams(), output_type="text", codec="utf-8"
        )
        return out.getvalue()
    except ImportError:
        raise ToneLoadError(
            "pdfminer.six is required to read PDF files. "
            "Run: pip install pdfminer.six"
        )
    except Exception as exc:
        raise ToneLoadError(f"PDF extraction failed: {exc}") from exc


# ── Helpers ───────────────────────────────────────────────────────────────────

def _get_ext(filename_lower: str) -> str:
    for ext in (".docx", ".doc", ".xlsx", ".xls", ".txt", ".csv", ".pdf", ".rtf"):
        if filename_lower.endswith(ext):
            return ext
    return ""


def _clean(text: str) -> str:
    """Normalise whitespace and strip control characters."""
    text = re.sub(r"[\r\n\t\x0b\x0c]+", " ", text)
    text = re.sub(r" {2,}", " ", text)
    # Remove common Word artefacts
    text = re.sub(r"[\x00-\x08\x0e-\x1f\x7f]", "", text)
    return text.strip()
